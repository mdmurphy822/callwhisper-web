"""
Transcript Export Service

Exports transcripts in multiple formats:
- JSON (structured with metadata)
- VTT (WebVTT subtitle format)
- CSV (tabular transcript data)
- PDF (formatted document)
- DOCX (Word document)
"""

from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
import json
import csv
import re

from ..core.logging_config import get_logger
from ..core.state import app_state

logger = get_logger(__name__)


class TranscriptExporter:
    """Export transcripts to various formats."""

    def __init__(self, output_folder: Path):
        self.output_folder = output_folder
        self.transcript_txt = output_folder / "transcript.txt"
        self.transcript_srt = output_folder / "transcript.srt"

    def _load_transcript_text(self) -> str:
        """Load plain text transcript."""
        if self.transcript_txt.exists():
            return self.transcript_txt.read_text(encoding="utf-8")
        return ""

    def _load_srt_entries(self) -> List[Dict[str, Any]]:
        """Parse SRT file into list of entries."""
        if not self.transcript_srt.exists():
            return []

        entries = []
        content = self.transcript_srt.read_text(encoding="utf-8")
        blocks = content.strip().split("\n\n")

        for block in blocks:
            lines = block.strip().split("\n")
            if len(lines) >= 3:
                # Parse timecode line: "00:00:00,000 --> 00:00:05,000"
                timecode = lines[1]
                match = re.match(
                    r"(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})", timecode
                )
                if match:
                    entries.append(
                        {
                            "index": int(lines[0]),
                            "start": match.group(1),
                            "end": match.group(2),
                            "text": "\n".join(lines[2:]),
                        }
                    )
        return entries

    def _get_metadata(self, recording_id: str) -> Dict[str, Any]:
        """Get recording metadata from app state."""
        recording = next(
            (r for r in app_state.completed_recordings if r.id == recording_id), None
        )
        if recording:
            return {
                "recording_id": recording.id,
                "ticket_id": recording.ticket_id,
                "created_at": recording.created_at.isoformat(),
                "duration_seconds": recording.duration_seconds,
            }
        return {"recording_id": recording_id}

    async def export_json(self, recording_id: str) -> Path:
        """
        Export as structured JSON with full metadata.

        Structure:
        {
            "version": "1.0.0",
            "generator": "CallWhisper",
            "exported_at": "ISO timestamp",
            "recording": { metadata },
            "transcript": {
                "text": "full text",
                "word_count": N,
                "segments": [ {start, end, text} ]
            }
        }
        """
        text = self._load_transcript_text()
        segments = self._load_srt_entries()
        metadata = self._get_metadata(recording_id)

        export_data = {
            "version": "1.0.0",
            "generator": "CallWhisper",
            "exported_at": datetime.now().isoformat(),
            "recording": metadata,
            "transcript": {
                "text": text,
                "word_count": len(text.split()) if text else 0,
                "segments": segments,
            },
        }

        output_path = self.output_folder / f"{recording_id}_transcript.json"
        output_path.write_text(
            json.dumps(export_data, indent=2, ensure_ascii=False), encoding="utf-8"
        )

        logger.info(
            "export_json_created", recording_id=recording_id, path=str(output_path)
        )
        return output_path

    async def export_vtt(self, recording_id: str) -> Path:
        """
        Export as WebVTT subtitle format.

        VTT format:
        WEBVTT

        00:00:00.000 --> 00:00:05.000
        Text here
        """
        entries = self._load_srt_entries()

        vtt_lines = ["WEBVTT", ""]
        for entry in entries:
            # Convert SRT timecode (00:00:00,000) to VTT (00:00:00.000)
            start = entry["start"].replace(",", ".")
            end = entry["end"].replace(",", ".")
            vtt_lines.append(f"{start} --> {end}")
            vtt_lines.append(entry["text"])
            vtt_lines.append("")

        output_path = self.output_folder / f"{recording_id}_transcript.vtt"
        output_path.write_text("\n".join(vtt_lines), encoding="utf-8")

        logger.info(
            "export_vtt_created",
            recording_id=recording_id,
            path=str(output_path),
            segments=len(entries),
        )
        return output_path

    async def export_csv(self, recording_id: str) -> Path:
        """
        Export as CSV with segment data.

        Columns: index, start_time, end_time, text
        """
        entries = self._load_srt_entries()

        output_path = self.output_folder / f"{recording_id}_transcript.csv"
        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["index", "start_time", "end_time", "text"])
            for entry in entries:
                writer.writerow(
                    [
                        entry["index"],
                        entry["start"],
                        entry["end"],
                        entry["text"].replace("\n", " "),
                    ]
                )

        logger.info(
            "export_csv_created",
            recording_id=recording_id,
            path=str(output_path),
            rows=len(entries),
        )
        return output_path

    async def export_pdf(self, recording_id: str) -> Path:
        """
        Export as formatted PDF document.

        Requires: reportlab
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        text = self._load_transcript_text()
        metadata = self._get_metadata(recording_id)

        output_path = self.output_folder / f"{recording_id}_transcript.pdf"

        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title", parent=styles["Heading1"], fontSize=18, spaceAfter=12
        )
        meta_style = ParagraphStyle(
            "Meta",
            parent=styles["Normal"],
            fontSize=10,
            textColor="#666666",
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "Body", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=12
        )

        story = []

        # Title
        story.append(Paragraph("Call Transcript", title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Metadata
        story.append(
            Paragraph(
                f"Recording ID: {metadata.get('recording_id', 'N/A')}", meta_style
            )
        )
        if metadata.get("ticket_id"):
            story.append(Paragraph(f"Ticket: {metadata['ticket_id']}", meta_style))
        if metadata.get("created_at"):
            story.append(Paragraph(f"Date: {metadata['created_at']}", meta_style))
        if metadata.get("duration_seconds"):
            mins = int(metadata["duration_seconds"] // 60)
            secs = int(metadata["duration_seconds"] % 60)
            story.append(Paragraph(f"Duration: {mins}:{secs:02d}", meta_style))

        story.append(Spacer(1, 0.3 * inch))

        # Transcript paragraphs
        for para in text.split("\n\n"):
            if para.strip():
                # Escape special characters for reportlab
                safe_para = (
                    para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                )
                story.append(Paragraph(safe_para, body_style))

        doc.build(story)

        logger.info(
            "export_pdf_created", recording_id=recording_id, path=str(output_path)
        )
        return output_path

    async def export_docx(self, recording_id: str) -> Path:
        """
        Export as Word document.

        Requires: python-docx
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        text = self._load_transcript_text()
        metadata = self._get_metadata(recording_id)

        doc = Document()

        # Title
        title = doc.add_heading("Call Transcript", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Metadata table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        def add_meta_row(label: str, value: str):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value

        add_meta_row("Recording ID", metadata.get("recording_id", "N/A"))
        if metadata.get("ticket_id"):
            add_meta_row("Ticket", metadata["ticket_id"])
        if metadata.get("created_at"):
            add_meta_row("Date", metadata["created_at"])
        if metadata.get("duration_seconds"):
            mins = int(metadata["duration_seconds"] // 60)
            secs = int(metadata["duration_seconds"] % 60)
            add_meta_row("Duration", f"{mins}:{secs:02d}")

        doc.add_paragraph()  # Spacer

        # Transcript heading
        doc.add_heading("Transcript", level=2)

        # Transcript text
        for para in text.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(12)

        output_path = self.output_folder / f"{recording_id}_transcript.docx"
        doc.save(str(output_path))

        logger.info(
            "export_docx_created", recording_id=recording_id, path=str(output_path)
        )
        return output_path


def get_exporter(output_folder: Path) -> TranscriptExporter:
    """Factory function for TranscriptExporter."""
    return TranscriptExporter(output_folder)
